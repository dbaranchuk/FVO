#-.- coding: utf-8 -.-
from app import app, db
from flask_login import current_user
from sqlalchemy import text
from sqlalchemy.sql import select, and_
from docx import Document as Doc

from openpyxl import load_workbook

from app.models import User, VUS, Document, Student_info, Basic_information, Comments
from app.models import Certificates_change_name, Communications, Passports, International_passports
from app.models import Registration_certificates, Middle_education, Spec_middle_education
from app.models import High_education, Military_education, Languages, Mothers_fathers
from app.models import Brothers_sisters_children, Married_certificates, Personal_data
from app.models.easy import *

from werkzeug.security import generate_password_hash
from flask import request, send_from_directory
import datetime
import json
from app.views.easy import *
from app.models.easy import *
from app.config import USER_PATH
import os
import random, string
import re
from keywords import *
from transliteration import *
from zipfile import ZipFile, ZIP_DEFLATED
import sys

def create_account(login, password, userData):
    hash = generate_password_hash(password)

    new_user = User(login = login, password = hash, entrance_year = int(userData['year']))
    new_user.vus_id = userData['vus'].id
    
    student_info = Student_info()

    new_user.students_info = student_info

    # add basic information
    basic_information = Basic_information(last_name=userData['lastName'], first_name=userData['firstName'], 
                                          middle_name=userData['middleName'])
    student_info['basic_information'] = basic_information
    db.session.add(basic_information)

    # add comments 
    comments = Comments()
    student_info['comments'] = comments
    db.session.add(comments)

    # common interface tables
    for table in get_user_tables():
        if table != 'basic_information':
            section = get_class_by_tablename(table)()
            if section.is_fixed:
                student_info[table] = section
                student_info['table_'+table] = TABLE_STATES['NOT_EDITED']
                db.session.add(section)
    
    db.session.add(new_user)
    db.session.add(student_info)
    db.session.commit()
    return True

def create_admin_account(data):
    hash = generate_password_hash(data['password'])

    new_user = User(login = data['login'], 
                    password = hash,
                    vus_id = int( data['vus_id'] ),
                    role = USER_STATES[ data['role'] ]
                    )

    db.session.add(new_user)
    db.session.commit()
    return True

@app.route('/comment_user', methods=['POST'])
def comment_user():
    data = request.form
    user_id = data['id'];
    comment = data['comment']
    comment_user = Comments.query.get(user_id);
    if(comment_user):
        comment_user.comment = unicode(comment);
    else:
        comment_user = Comments(
            id = user_id,
            comment = unicode(comment)
        )
    db.session.add(comment_user)
    db.session.commit()
    return gen_success()

@app.route('/make_account', methods=['POST'])
def make_account():
    data = request.form
    if 'login' not in data or 'password' not in data:
        return gen_error('Wrong data sent to server (must be [login, password]).')
    user = User.query.filter_by(login=data['login']).first()
    if user:
        return gen_success(message={'status':'error', 'error' : u'Пользователь с таким логином уже существует'})
    create_admin_account(data)
    return gen_success(message={'status':'ok'})

@app.route('/post_add_vus', methods=['POST'])
def post_add_vus():
    data = request.form
    vus = VUS.query.filter_by(number=data['number'], code=data['code']).first()
    if vus:
        return gen_success(message={'status':'error', 'error' : u'Специальность была добавлена ранее'})
    vus = VUS()
    vus.number=data['number']
    vus.code=data['code']
    vus.name1=data['name1']
    vus.name2=data['name2']
    db.session.add(vus)
    db.session.commit()
    return gen_success(message={'status':'ok'})

@app.route('/create_accounts', methods=['POST'])
def create_accounts():
    if 'file' not in request.files:
        return gen_error('Файл не выбран')
    file = request.files['file']

    # if user does not select file, browser also
    # submit a empty part without filename
    if file.filename == '' or not file:
        return gen_error(u'Файл не выбран')
    if file.filename[-4:] != 'xlsx':
        return gen_error(u'Файл должен быть формата .xlsx')

    if 'vus' not in request.form:
        return gen_error('Выберите ВУС')
    if 'completionYear' not in request.form:
        return gen_error('Введите год поступления')

    completionYear = request.form['completionYear']
    if completionYear == '':
        return gen_error('Введите год поступления')

    vus = map(int, request.form['vus'].split())
    vus = VUS.query.filter_by(number=vus[0], code=vus[1]).first()
    if vus is None:
        return gen_error('Such vus not yet exists in this system')

    wb = load_workbook(file)
    active = wb.active
    userNames = User.query.with_entities(User.login)

    for idx, row in enumerate(active.rows, start = 1):
        login = ''

        #фамилия
        for char in row[0].value:
            login += vocabulary[char.lower()]
        login += u'.'

        #инициалы имени и отчества
        firstNameShort = row[1].value[0].lower()
        middleNameShort = row[2].value[0].lower()
        login += vocabulary[firstNameShort] + u'.'
        login += vocabulary[middleNameShort] + u'.'
        
        login += completionYear
        password = ''.join(random.choice(string.ascii_uppercase + string.ascii_lowercase + string.digits) for _ in range(8))

        for name in userNames:
            if login == name.login:
                return gen_error(u'В системе уже существует аккаунт: ' + login)

        info = {
            'lastName': row[0].value,
            'firstName': row[1].value,
            'middleName': row[2].value,
            'year': completionYear,
            'vus': vus
        }

        create_account(login, password, info)

        active.cell(row = idx, column = 4, value = login)
        active.cell(row = idx, column = 5, value = password)

    path = os.path.join(USER_PATH, 'logins.xlsx')
    wb.save(path)

    return gen_success(url = '/static/user_data/logins.xlsx')


@app.route('/add_document', methods=['POST'])
def add_document():
    if 'file' not in request.files:
        return gen_error('No file sent')
    file = request.files['file']

    # if user does not select file, browser also
    # submit a empty part without filename
    if file.filename == '' or not file:
        return gen_error('No file selected')

    if 'name' not in request.form:
        return gen_error('No name document could not be created')
    if 'vus' not in request.form:
        return gen_error('No vus document could not be created')

    vus = map(int, request.form['vus'].split())
    vus = VUS.query.filter_by(number=vus[0], code=vus[1]).first()
    if vus is None:
        return gen_error('Such vus not yet exists in this system')

    docs = os.listdir(os.path.join(USER_PATH, 'documents'))
    filename = file.filename
    if filename in docs:
        i = 1
        while filename + str(i) in docs:
            i += 1
        filename += str(i)

    file.save(os.path.join(USER_PATH, 'documents', filename))
    d = Document(name=request.form['name'], vus_id=vus.id, filename=filename)
    db.session.add(d)
    db.session.commit()

    return gen_success(filename=filename, message='Success!')

@app.route('/delete_document', methods=['POST'])
def delete_document():
    data = json.loads(request.data)
    docId = data['docId']

    document = Document.query.filter_by(id = docId).first()
    filePath = os.path.join(USER_PATH, 'documents', document.filename)


    if document is None:
        return gen_error('No document with such id')

    os.remove(filePath)

    Document.query.filter_by(id = docId).delete()
    db.session.commit()

    return gen_success(message='Success!')

### POSTs

def check_errors_in_input_data(table, data):
    tableclass = get_class_by_tablename(table)()
    errors = []
    for field, value in data.iteritems():
        if len(field) and not len(value):
            errors.append( u'Заполните поле "' + tableclass.get_russian_name( field ) + '"' )
    return errors

def save_not_fixed_section_information(data):
    elements = json.loads(data['elements'])
    # проверяем пустые поля
    errors = []
    for element in elements:
        errors += check_errors_in_input_data(data['table'], element)

    if len(errors):
        return gen_success(message = {'status':'error', 'errors':"<br>".join(errors)})

    student_info = User.query.get( current_user.id ).students_info
    records = student_info[data['table']]
    tableclass = get_class_by_tablename(data['table'])
    
    ### delete existing
    if records and len(records):
        for record in records:
            db.session.delete(record)

    ### add new records
    for element in elements:
        element_fields = { field : element[field] for field in element if hasattr(tableclass, field) }
        new_record = tableclass()
        for field, value in element_fields.iteritems():
            new_record[field] = value
        new_record.student_info    = student_info
        new_record.student_info_id = student_info.id
        db.session.add(new_record)
        db.session.commit()

    student_info['table_' + data['table']] = TABLE_STATES['EDITED']
    db.session.commit()

    return gen_success(message = {'status':'ok'} )
        

def save_section_information(data):
    
    # проверяем пустые поля
    if 'elements' in data:
        return save_not_fixed_section_information(data)

    errors = check_errors_in_input_data(data['table'], data)
    if len(errors):
        return gen_success(message = {'status':'error', 'errors' : "<br>".join(errors) })
    
    student_info = User.query.get( current_user.id ).students_info
    table = student_info[data['table']]
    # обновляем таблицу
    if table:
        tableclass = get_class_by_tablename(data['table'])
        table_fields = { field : data[field] for field in data if hasattr(tableclass, field) }
        for field, value in table_fields.iteritems():
            table[field] = value
        student_info['table_' + data['table']] = TABLE_STATES['EDITED']
        db.session.commit()

    return gen_success(message = {'status':'ok'} )

def approve_all_sections(data):
    user = User.query.get( int(data['user_id']) )
    student_info = user.students_info
    for table in get_user_tables():
        student_info['table_' + table] = TABLE_STATES['APPROVED']
    user.approved = True
    db.session.commit()
    return gen_success(message = {'status':'ok'} )

def change_section_state(data):
                    
    new_state = int(data['new_state'])
    user = User.query.get( int(data['user_id']) )
    student_info = user.students_info

    student_info['table_' + data['table']] = new_state
    if new_state == TABLE_STATES['DECLINED']:
        user.approved = False
        student_info['comments'][data['table'] + '_comment'] = data['comment']
    else:
        student_info['comments'][data['table'] + '_comment'] = ''

    is_all_approved = True
    for table in get_user_tables():
        if student_info['table_' + table] != TABLE_STATES['APPROVED']:
            is_all_approved = False
            break

    if is_all_approved:
        user.approved = True

    db.session.commit()
    return gen_success(message = {'status':'ok'} )

def send_quiz_to_check(data):
    return gen_success(message =  {'status':'ok'})

### SEARCH
def searchUsers(data):

    sqlRequest = getSqlRequest(data['lastName'], data['year'], data['vus'])

    requestResult = db.engine.execute(sqlRequest)

    searchResult = []
    for row in requestResult:
        matchedUser = {
            'id' : row[0],
            'lastName' : row[3],
            'year' : row[2][-4:],
            'vus' : '%03d %03d' % (row[4], row[5])
        }
        searchResult.append(matchedUser)

    return gen_success(result = searchResult)

def getSqlRequest(lastName, year, vusStr):

    sqlRequest = "select u_id, u_role, u_login, bi_last_name,\
        VUS.number as 'vus_num',\
        VUS.code as 'vus_code'\
        from (\
        select user.id as 'u_id',\
        user.role as 'u_role',\
        user.login as 'u_login',\
        user.vus_id as 'u_vus_id',\
        bi_last_name\
        from (\
        select\
        student_info.user_id as 'si_user_id',\
        basic_information.last_name as 'bi_last_name'\
        from student_info left join basic_information\
        on student_info.id = basic_information.student_info_id) as X\
        left join user\
        on X.si_user_id = user.id) as Y\
        left join VUS\
        on Y.u_vus_id = VUS.id "

    conds = ["u_role = '0'"]

    if lastName != '':
        conds.append("bi_last_name = '" + unicode(lastName) + "'")

    if year != '':
        conds.append("u_login like '%" + unicode(year) + "'")

    if vusStr != u'не выбрано':
        vus = map(int, vusStr.split())

        conds.append("vus_num = '" + str(vus[0]) + "'")
        conds.append("vus_code = '" + str(vus[1]) + "'")


    whereBlock = 'where ' + ' and '.join(conds) + ';'

    return text(sqlRequest + whereBlock)

### GENERATING DOCUMENTS

def generateDocuments(data):
    userIDs = json.loads(data['userIDs'])
    docIDs = json.loads(data['docIDs'])
    
    users = User.query.filter(User.id.in_(userIDs)).all()
    documents = Document.query.filter(Document.id.in_(docIDs)).all()

    if not users:
        return gen_success(success = False, message = 'Выберите хотя бы одного пользователя')
    if not documents:
        return gen_success(success = False, message = 'Выберите хотя бы один документ')
    
    for user in users:
        accessor = Students_info_lables_accessor(user.students_info)
        for document in documents:
            docPath = os.path.join(USER_PATH, 'documents', document.filename)
            doc = Doc(docPath)
            
            regex = re.compile('\{[a-zA-Z0-9_.@]+\}')
            
            for p in doc.paragraphs:
                iterator = regex.finditer(p.text)
                for match in iterator:
                    keyword = match.group()
                    value = unicode(accessor[keyword]) if unicode(accessor[keyword]) != None else u'НЕПРАВИЛЬНЫЙ КЛЮЧ!'
                    text = p.text.replace(keyword, value)
                    style = p.style
                    p.text = text
                    p.style = style
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            iterator = regex.finditer(p.text)
                            for match in iterator:
                                keyword = match.group()
                                value = unicode(accessor[keyword]) if unicode(accessor[keyword]) != None else u'НЕПРАВИЛЬНЫЙ КЛЮЧ!'
                                text = p.text.replace(keyword, value)
                                style = p.style
                                p.text = text
                                p.style = style
                    
            doc_name = os.path.join(USER_PATH, 'documents', 'temp', document.filename[:-5] + str(user.id) + '.docx')
            doc.save(doc_name)
    
    zippath = os.path.join(USER_PATH, 'Documents.zip')
    zipf = ZipFile(zippath, 'w', ZIP_DEFLATED)
    
    basedir = os.path.join(USER_PATH, 'documents', 'temp')
    for root, dirs, files in os.walk(basedir):
        for fn in files:
            absfn = os.path.join(root, fn)
            zfn = absfn[len(basedir)+len(os.sep):]
            zipf.write(absfn, zfn)
            os.remove(absfn)
    zipf.close()

    return gen_success(url = '/static/user_data/Documents.zip', success = True)



@app.route('/post_query', methods=['POST'])
def post_query():
    data = request.form
    try:
        if 'do' in data and data['do'] in POST_METHODS:
            return POST_METHODS[data['do']](data)
        else:
            return gen_success(message = {'status':'error', 'error':'post method not defined'})
    except Exception as err:
        return gen_success(message = {'status':'error', 'error':'error in post method'})


# list of post methods
POST_METHODS = dict( [ (table, save_section_information) for table in get_user_tables() ] )
POST_METHODS.update( {
                'searchUsers': searchUsers,
                'generateDocuments': generateDocuments,
                'send_quiz_to_check': send_quiz_to_check,
                'change_section_state' : change_section_state,
                'approve_all_sections' : approve_all_sections,

                })


